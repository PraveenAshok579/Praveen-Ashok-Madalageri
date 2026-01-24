# ======================================================
# CBR PREDICTION WEB APP – FINAL Q1 JOURNAL VERSION
# (LANDFILL + PAVEMENT SUBGRADE APPLICATIONS)
# ======================================================

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from datetime import datetime
import io

from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.ensemble import RandomForestRegressor

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

# ------------------------------------------------------
# PAGE CONFIGURATION
# ------------------------------------------------------
st.set_page_config(
    page_title="CBR prediction using machine learning models",
    page_icon="🚦",
    layout="centered"
)

# ------------------------------------------------------
# DARK THEME + WHITE TEXT
# ------------------------------------------------------
st.markdown("""
<style>
.stApp {
    background-color: #0e1117;
    color: white;
}
h1, h2, h3, p {
    color: white !important;
    text-align: center;
}
div[data-testid="stAlert"] p {
    color: white !important;
}
button {
    background-color: #00c853 !important;
    color: black !important;
    font-size: 18px !important;
    border-radius: 8px !important;
}
</style>
""", unsafe_allow_html=True)

# ------------------------------------------------------
# TITLE
# ------------------------------------------------------
st.markdown("<h1>🚦 CBR prediction using machine learning models</h1>", unsafe_allow_html=True)
st.markdown(
    "<p style='color:#b0b0b0;'>"
    "Decision-support tool for landfill liner & cover and pavement subgrade applications"
    "</p>",
    unsafe_allow_html=True
)

# ------------------------------------------------------
# LOAD DATA
# ------------------------------------------------------
@st.cache_data
def load_data():
    return pd.read_excel("CBR Data.xlsx")

df = load_data()

TARGET = "CBR"
X = df.drop(columns=[TARGET])
y = df[TARGET]

# ------------------------------------------------------
# MODEL TRAINING
# ------------------------------------------------------
X_train, _, y_train, _ = train_test_split(
    X, y, test_size=0.2, random_state=42
)

model = Pipeline([
    ("scaler", StandardScaler()),
    ("rf", RandomForestRegressor(
        n_estimators=300,
        max_depth=15,
        min_samples_split=4,
        random_state=42
    ))
])

model.fit(X_train, y_train)

# ------------------------------------------------------
# INPUT SECTION
# ------------------------------------------------------
st.subheader("🔢 Soil Input Parameters")

colors = ["#ff5252", "#ffa000", "#ffee58", "#69f0ae", "#40c4ff", "#b388ff"]
input_data = {}
cols = st.columns(2)

for i, col in enumerate(X.columns):
    with cols[i % 2]:
        st.markdown(
            f"<div style='border-left:5px solid {colors[i % len(colors)]}; "
            f"padding-left:10px;margin-bottom:5px;'>"
            f"<strong style='color:{colors[i % len(colors)]}'>{col}</strong>"
            f"</div>",
            unsafe_allow_html=True
        )
        input_data[col] = st.number_input(
            "",
            min_value=float(df[col].min()),
            max_value=float(df[col].max()),
            value=float(df[col].mean()),
            key=col
        )

input_df = pd.DataFrame([input_data])

# ------------------------------------------------------
# PREDICTION
# ------------------------------------------------------
if st.button("🔍 Predict CBR"):

    prediction = float(model.predict(input_df)[0])

    if prediction < 3:
        color = "red"
        quality = "Very Poor"
    elif prediction < 5:
        color = "orange"
        quality = "Poor"
    elif prediction < 10:
        color = "yellow"
        quality = "Fair"
    else:
        color = "green"
        quality = "Good"

    # Gauge
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=prediction,
        number={"suffix": " %"},
        gauge={
            "axis": {"range": [0, 20]},
            "bar": {"color": color},
            "steps": [
                {"range": [0, 3], "color": "red"},
                {"range": [3, 5], "color": "orange"},
                {"range": [5, 10], "color": "yellow"},
                {"range": [10, 20], "color": "green"}
            ]
        },
        title={"text": "Predicted CBR"}
    ))

    st.plotly_chart(fig, use_container_width=True)
    st.success(f"**Subgrade Quality:** {quality}")

    # ==================================================
    # LANDFILL CONCLUSION
    # ==================================================
    st.subheader("🧾 Engineering Conclusion – Landfill Liner & Cover")

    if prediction < 3:
        landfill_conclusion = (
            "The predicted CBR indicates a very weak soil condition. "
            "Such soils are not suitable for direct use in landfill liner or cover systems "
            "due to inadequate load-bearing capacity and poor constructability. "
            "Soil stabilization, blending, or replacement is mandatory in accordance with "
            "EPA and CPCB landfill engineering practice."
        )
    elif prediction < 5:
        landfill_conclusion = (
            "The predicted CBR represents marginal strength. "
            "The soil may be conditionally used for landfill cover applications after "
            "appropriate improvement measures. Direct application as a landfill liner "
            "is not recommended."
        )
    elif prediction < 10:
        landfill_conclusion = (
            "The predicted CBR indicates moderate strength. "
            "The soil is suitable for landfill cover systems and may be considered for "
            "liner applications with controlled compaction and quality assurance measures."
        )
    else:
        landfill_conclusion = (
            "The predicted CBR indicates good strength and stability. "
            "The soil is suitable for both landfill liner and cover applications, "
            "ensuring adequate resistance to deformation during construction and operation."
        )

    st.markdown(
        f"<div style='color:white; background-color:#1e222a; "
        f"padding:15px; border-radius:10px;'>{landfill_conclusion}</div>",
        unsafe_allow_html=True
    )

    # ==================================================
    # PAVEMENT SUBGRADE CONCLUSION
    # ==================================================
    st.subheader("🛣️ Engineering Conclusion – Pavement Subgrade")

    if prediction < 3:
        pavement_conclusion = (
            "The predicted CBR indicates a very poor pavement subgrade. "
            "Such soils are unsuitable for direct pavement construction and "
            "require stabilization, replacement, or provision of a thick capping layer "
            "to prevent excessive deformation and premature pavement failure."
        )
    elif prediction < 5:
        pavement_conclusion = (
            "The predicted CBR represents a poor subgrade condition. "
            "The soil may be used for pavement construction only after significant "
            "improvement measures such as lime or cement stabilization."
        )
    elif prediction < 10:
        pavement_conclusion = (
            "The predicted CBR indicates a fair subgrade condition. "
            "The soil is suitable for low to medium traffic pavements provided "
            "adequate pavement thickness is designed."
        )
    else:
        pavement_conclusion = (
            "The predicted CBR indicates good subgrade strength. "
            "The soil is suitable for flexible pavement construction with "
            "conventional pavement layer thicknesses."
        )

    st.markdown(
        f"<div style='color:white; background-color:#1e222a; "
        f"padding:15px; border-radius:10px;'>{pavement_conclusion}</div>",
        unsafe_allow_html=True
    )

    # ==================================================
    # REPORT GENERATION
    # ==================================================
    now = datetime.now().strftime("%d-%m-%Y %H:%M")

    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=A4)
    c.setFont("Times-Roman", 12)

    c.drawString(50, 800, "CBR Prediction Report")
    c.drawString(50, 780, f"Generated on: {now}")
    c.drawString(50, 760, f"Predicted CBR: {prediction:.2f} %")

    c.drawString(50, 730, "Landfill Application Conclusion:")
    c.drawString(50, 710, landfill_conclusion[:100])

    c.drawString(50, 670, "Pavement Subgrade Conclusion:")
    c.drawString(50, 650, pavement_conclusion[:100])

    c.save()
    pdf_buffer.seek(0)

    doc = Document()
    doc.add_heading("CBR Prediction Report", level=1)
    doc.add_paragraph(f"Generated on: {now}")
    doc.add_paragraph(f"Predicted CBR: {prediction:.2f} %")

    doc.add_heading("Landfill Application Conclusion", level=2)
    doc.add_paragraph(landfill_conclusion)

    doc.add_heading("Pavement Subgrade Conclusion", level=2)
    doc.add_paragraph(pavement_conclusion)

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    col1, col2 = st.columns(2)
    col1.download_button("📄 Download PDF Report", pdf_buffer, "CBR_Report.pdf", "application/pdf")
    col2.download_button("📝 Download Word Report", word_buffer, "CBR_Report.docx")

# ------------------------------------------------------
# FOOTER
# ------------------------------------------------------
st.markdown("---")
st.markdown(
    "<p style='text-align:center;color:#888;'>"
    "Decision-support tool for environmental and transportation geotechnics"
    "</p>",
    unsafe_allow_html=True
)
